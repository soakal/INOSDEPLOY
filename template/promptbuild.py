#!/usr/bin/env python3
"""
promptbuild.py — SOP Factory image reference pack builder
==========================================================

This script is a HELPER ONLY. It does NOT write procedure text or SOP content.

What it does:
  - Reads numbered images from active\\ (produced by prep_active.py)
  - Produces a clean Word document listing every image with its caption
  - That document is a REFERENCE for Claude when it writes the real SOP

The real SOP is written by Claude in a separate step. See PROMPT.txt for
the full workflow.

Usage:
  python promptbuild.py              # auto-locate from script folder
  python promptbuild.py --dry-run   # check paths without writing
"""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _resolve_paths(prompt_path: Path) -> dict:
    this_folder = prompt_path.parent
    root = this_folder.parent
    active_dir = root / "active"
    output_dir = root / "output"

    # Parse Section [3] output filename from PROMPT.txt
    text = _read_text(prompt_path)
    out_path = None
    in_s3 = False
    for line in text.splitlines():
        s = line.strip()
        if s.startswith("[3]"):
            in_s3 = True
            continue
        if in_s3 and re.match(r"^\[\d\]", s):
            in_s3 = False
        if in_s3 and s.lower().endswith(".docx"):
            candidate = s.replace("[this folder]", str(this_folder)).replace("/", "\\")
            out_path = Path(candidate).expanduser()
            break

    if out_path is None:
        out_path = output_dir / "SOP_IMAGE_PACK.docx"

    out_str = str(out_path).replace("\\_output\\", "\\output\\")
    out_path = Path(out_str)
    return {"active_dir": active_dir, "output_path": out_path}


def _load_manifest(active_dir: Path) -> dict[str, str]:
    manifest = active_dir / "image_manifest.txt"
    if not manifest.exists():
        return {}
    mapping: dict[str, str] = {}
    lines = _read_text(manifest).splitlines()
    i = 0
    while i < len(lines):
        name = lines[i].strip()
        if re.match(r"^\d+_.+\.(png|jpg|jpeg|bmp|gif|webp)$", name, flags=re.I):
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            desc = ""
            if j < len(lines):
                m = re.match(r"^\s*[→\-]+?\s*(.*)$", lines[j].strip())
                desc = m.group(1).strip() if m else lines[j].strip()
            mapping[name] = desc.strip("[]") or "Image"
            i = j + 1
            continue
        i += 1
    return mapping


def _numbered_images(active_dir: Path) -> list[Path]:
    imgs = [
        p for p in active_dir.iterdir()
        if p.is_file() and re.match(r"^\d+_.+\.(png|jpg|jpeg|bmp|gif|webp)$", p.name, flags=re.I)
    ]
    imgs.sort(key=lambda p: p.name)
    return imgs


def build_image_pack(prompt_path: Path, *, dry_run: bool = False) -> Path:
    """
    Build a clean image reference pack: one image per page, with caption.
    This is NOT the final SOP — it is a reference for Claude to use.
    """
    resolved = _resolve_paths(prompt_path)
    active_dir: Path = resolved["active_dir"]
    output_path: Path = resolved["output_path"]

    if not active_dir.exists():
        raise FileNotFoundError(f"active\\ not found: {active_dir}")

    images = _numbered_images(active_dir)
    manifest_map = _load_manifest(active_dir)

    print(f"Images found : {len(images)}")
    print(f"Output path  : {output_path}")

    if not images:
        print("No numbered images found. Run prep_active.py first.")
        return output_path

    if dry_run:
        for img in images:
            cap = manifest_map.get(img.name, "Image")
            print(f"  {img.name}  →  {cap}")
        return output_path

    output_path.parent.mkdir(parents=True, exist_ok=True)

    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Simple heading
    h = doc.add_heading("SOP Image Pack", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph(
        "Reference images for Claude SOP generation. "
        "Images are in the order they appear in active\\."
    )
    note.runs[0].italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    inserted = 0
    for img in images:
        caption = manifest_map.get(img.name, img.stem)

        # image paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(str(img), width=Inches(6.0))
        except Exception as e:
            p.add_run(f"[Could not insert {img.name}: {e}]")

        # caption
        cap_p = doc.add_paragraph(f"{img.name}  —  {caption}")
        cap_p.runs[0].italic = True
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        inserted += 1

    doc.save(str(output_path))
    print(f"Inserted      : {inserted} images")
    print(f"Saved         : {output_path}")
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Build SOP image reference pack.")
    parser.add_argument("--prompt", help="Path to PROMPT.txt (auto-detected if omitted)")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    prompt_path = (
        Path(args.prompt).resolve()
        if args.prompt
        else Path(__file__).resolve().parent / "PROMPT.txt"
    )

    build_image_pack(prompt_path, dry_run=args.dry_run)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
